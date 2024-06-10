from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Himal Sharma', 0)

# Contact Information
doc.add_paragraph('Software Developer\nHorsens, Denmark\nhimal28925@gmail.com | [GitHub](https://github.com/himal28924) | [LinkedIn](https://www.linkedin.com/in/himal-sharma-673937144/)')

# Add a line break
doc.add_paragraph('')

# Professional Summary
doc.add_heading('Professional Summary:', level=1)
doc.add_paragraph(
    'Software engineering student skilled in applying modern technologies, algorithms, and Agile methodologies to develop efficient, customer-centric solutions. '
    'Experienced in collaborative problem-solving and delivering high-quality software through teamwork.'
)

# Add a line break
doc.add_paragraph('')

# Professional Experience
doc.add_heading('Professional Experience:', level=1)

experience = doc.add_paragraph()
experience.add_run('Software Engineering Intern\n').bold = True
experience.add_run('Beumer Group, Aarhus, Denmark\n').italic = True
experience.add_run('Feb 2023 – Jun 2023\n')

experience.add_paragraph(
    '- Developed a WPF application to centralize log management, reducing troubleshooting time by 30%.\n'
    '- Led Java-based projects to address critical system issues, enhancing customer support with new services.\n'
    '- Engineered and launched a SERIAL connection service, solving complex problems through extensive research.\n'
    '- Collaborated with customers to refine software development using scalable solutions and modern practices.\n'
    '- Employed Agile methodologies to expedite software delivery and nurture continuous improvement.'
)

# Add a line break
doc.add_paragraph('')

# Education
doc.add_heading('Education:', level=1)

education = doc.add_paragraph()
education.add_run('Bachelor of Engineering in Software Technology\n').bold = True
education.add_run('VIA University College, Horsens, Denmark\n').italic = True
education.add_run('Feb 2022 – Aug 2024\n')

education.add_paragraph(
    '- Practical experience in IoT development, web development, and distributed system design.\n'
    '- Applied Agile methodologies and fundamental design principles in software development projects.\n'
    '- Deepened understanding of networking, securities, and data management.'
)

education.add_paragraph('')

education.add_run('GCE A-Level\n').bold = True
education.add_run('British Model College, Kathmandu, Nepal\n').italic = True
education.add_run('2018\n')

# Add a line break
doc.add_paragraph('')

# Projects
doc.add_heading('Projects:', level=1)

projects = doc.add_paragraph()
projects.add_run('Bachelor Project: ').bold = True
projects.add_run('Web Application using .NET and EFC for backend and React for frontend\n')
projects.add_paragraph(
    '- Developed a web application implementing domain-driven design principles.\n'
    '- Conducted unit and integration testing, and utilized CI/CD pipelines for automated testing and deployment.'
)

projects.add_paragraph('')

projects.add_run('Best Movies:\n').bold = True
projects.add_paragraph(
    '- Developed a React front end and C# backend.\n'
    '- Leveraged Microsoft Azure for hosting, ensuring high availability and scalability.'
)

projects.add_paragraph('')

projects.add_run('Movie Recommendation:\n').bold = True
projects.add_paragraph(
    '- Engineered a machine learning model to predict film ratings, improving accuracy through hybrid model experimentation.'
)

projects.add_paragraph('')

projects.add_run('Remote Greenhouse Monitoring System:\n').bold = True
projects.add_paragraph(
    '- Developed an Arduino-based system for environmental tracking, using LoRaWAN for data transmission.'
)

projects.add_paragraph('')

projects.add_run('SEP3- Market Place:\n').bold = True
projects.add_paragraph(
    '- Constructed a three-tier system using Blazor, Java (SpringBoot), and ASP.NET, with REST API and gRPC services.'
)

# Add a line break
doc.add_paragraph('')

# Skills
doc.add_heading('Skills:', level=1)

skills = doc.add_paragraph()
skills.add_run('Programming Languages:\n').bold = True
skills.add_paragraph('- Java (Spring Boot), C# (.NET WPF Blazor), C (Embedded RTP), Python, F#')

skills.add_paragraph('')

skills.add_run('Web Development:\n').bold = True
skills.add_paragraph('- JavaScript, React, Blazor')

skills.add_paragraph('')

skills.add_run('Design Patterns & Principles:\n').bold = True
skills.add_paragraph('- MVVM, Observer Pattern, SOLID')

skills.add_paragraph('')

skills.add_run('Project Management:\n').bold = True
skills.add_paragraph('- Agile Methodology, Scrum')

skills.add_paragraph('')

skills.add_run('AI & Data:\n').bold = True
skills.add_paragraph('- SQL, EFC, Machine Learning, Neural Networks')

skills.add_paragraph('')

skills.add_run('Web Services:\n').bold = True
skills.add_paragraph('- REST, gRPC, Socket')

skills.add_paragraph('')

skills.add_run('Cloud & Testing:\n').bold = True
skills.add_paragraph('- Unit Testing, Google Test, Git, Cloud Computing, DevOps tools, CI/CD Pipeline')

# Add a line break
doc.add_paragraph('')

# Languages
doc.add_heading('Languages:', level=1)

languages = doc.add_paragraph()
languages.add_run('English (Fluent)\n').bold = True
languages.add_run('Hindi (Fluent)\n').bold = True
languages.add_run('Nepali (Fluent)\n').bold = True
languages.add_run('Dansk (Intermediate)\n').bold = True

# Add a line break
doc.add_paragraph('')

# Interests
doc.add_heading('Interests:', level=1)

interests = doc.add_paragraph()
interests.add_run('Technology innovation, continuous learning, sports')

# Save the document
doc_path = "/mnt/data/Himal_Sharma_Resume.docx"
doc.save(doc_path)

doc_path
